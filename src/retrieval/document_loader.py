"""
Document loading and processing utilities for RAG.

This module provides functions for loading and processing documents from 
various formats (including .docx) for use in retrieval-augmented generation.
"""

import os
from typing import List, Optional

from docx import Document as DocxDocument
from llama_index.core import Document as LlamaDocument
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader

import sys
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from config import DATA_DIR


class DocumentProcessor:
    """
    Handles document loading and processing for RAG applications.
    
    This class provides methods to load, parse, and index documents from
    various sources, particularly focusing on text and docx formats.
    """
    
    def __init__(self, data_dir: str = DATA_DIR, chunk_size: int = 1024, chunk_overlap: int = 20):
        """
        Initialize the document processor.
        
        Args:
            data_dir (str): Directory containing the documents
            chunk_size (int): Size of text chunks for indexing
            chunk_overlap (int): Overlap between consecutive chunks
        """
        self.data_dir = data_dir
        self.node_parser = SentenceSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap
        )
    
    def load_docx(self, file_path: str) -> str:
        """
        Load text content from a .docx file.
        
        Args:
            file_path (str): Path to the .docx file
            
        Returns:
            str: Extracted text content
        """
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def load_directory(self, directory: Optional[str] = None) -> List[LlamaDocument]:
        """
        Load all documents from a directory.
        
        Args:
            directory (str, optional): Directory to load from. Defaults to self.data_dir.
            
        Returns:
            List[LlamaDocument]: List of loaded documents
        """
        directory = directory or self.data_dir
        return SimpleDirectoryReader(directory).load_data()
    
    def load_single_document(self, file_path: str) -> LlamaDocument:
        """
        Load a single document from a file path.
        
        Args:
            file_path (str): Path to the document
            
        Returns:
            LlamaDocument: Loaded document
        """
        if file_path.endswith('.docx'):
            content = self.load_docx(file_path)
            return LlamaDocument(text=content)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return LlamaDocument(text=content)
    
    def process_documents(self, documents: List[LlamaDocument]) -> List:
        """
        Process documents into nodes for indexing.
        
        Args:
            documents (List[LlamaDocument]): Documents to process
            
        Returns:
            List: Processed nodes
        """
        return self.node_parser.get_nodes_from_documents(documents)
    
    def create_index(self, nodes=None, documents=None):
        """
        Create a vector store index from nodes or documents.
        
        Args:
            nodes: Pre-processed nodes (optional)
            documents: Documents to process (optional)
            
        Returns:
            VectorStoreIndex: The created index
        """
        if nodes is None and documents is None:
            documents = self.load_directory()
            nodes = self.process_documents(documents)
        elif nodes is None and documents is not None:
            nodes = self.process_documents(documents)
            
        return VectorStoreIndex(nodes)
    
    def get_query_engine(self, index=None):
        """
        Get a query engine from an index.
        
        Args:
            index: The index to query (optional)
            
        Returns:
            QueryEngine: A query engine for the index
        """
        if index is None:
            index = self.create_index()
            
        return index.as_query_engine() 