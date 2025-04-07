from llama_index.core.llama_dataset.rag import LabelledRagDataset
from llama_index.core.llama_dataset.generator import RagDatasetGenerator
from llama_index.core.evaluation import AnswerRelevancyEvaluator
from llama_index.core.node_parser import SentenceSplitter
import os
from typing import Any
from llama_index.core import Document as LlamaDocument
from llama_index.core.llms import (
    CustomLLM,
    CompletionResponse,
    CompletionResponseGen,
    LLMMetadata,
)
from llama_index.core.llms.callbacks import llm_completion_callback
from transformers import AutoTokenizer, AutoModelForCausalLM
from llama_index.core import Settings, VectorStoreIndex, SimpleDirectoryReader
from datetime import datetime
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core.evaluation import ContextRelevancyEvaluator
from docx import Document
from peft import PeftModel

class CustomLLM(CustomLLM):
    context_window: int = 128000  # 上下文窗口大小
    num_output: int = 18000  # 输出的token数量
    model_name: str = "./Qwen_QLoRA"  # 模型名称
    tokenizer: object = None  # 分词器
    model: object = None  # 模型
    dummy_response: str = "My response"
 
    def __init__(self, pretrained_model_name_or_path, finetuned_model_path, offload_dir):
        super().__init__()
 
        # GPU方式加载模型
        self.tokenizer = AutoTokenizer.from_pretrained(pretrained_model_name_or_path, device_map="cuda",
                                                       trust_remote_code=True)
        base_model = AutoModelForCausalLM.from_pretrained(
            pretrained_model_name_or_path,
            device_map="cuda",  
            trust_remote_code=True  # 允许加载自定义代码
        )
        self.model = PeftModel.from_pretrained(
            base_model,
            finetuned_model_path,  # 训练后的模型路径
            device_map="cuda",  # 指定设备为GPU
            offload_dir=offload_dir  # 添加卸载目录
        ).eval()
        self.model = self.model.float()
 
    @property
    def metadata(self) -> LLMMetadata:
        """Get LLM metadata."""
        # 得到LLM的元数据
        return LLMMetadata(
            context_window=self.context_window,
            num_output=self.num_output,
            model_name=self.model_name,
        )
 
 
    @llm_completion_callback()  # 回调函数
    def complete(self, prompt: str, **kwargs: Any) -> CompletionResponse:
        now = datetime.now()
        inputs = self.tokenizer.encode(prompt, return_tensors='pt').cuda()  # GPU方式
        # inputs = self.tokenizer.encode(prompt, return_tensors='pt')  # CPU方式
        outputs = self.model.generate(inputs, max_length=self.num_output)
        response = self.tokenizer.decode(outputs[0])
        # 完成函数
        # print(f"完成函数 当前时间为{now} response={response}")
        return CompletionResponse(text=response)
 
    @llm_completion_callback()
    def stream_complete(
            self, prompt: str, **kwargs: Any
    ) -> CompletionResponseGen:
        # 流式完成函数
        now = datetime.now()
        print("流式完成函数")
 
        inputs = self.tokenizer.encode(prompt, return_tensors='pt').cuda()  # GPU方式
        # inputs = self.tokenizer.encode(prompt, return_tensors='pt')  # CPU方式
        outputs = self.model.generate(inputs, max_length=self.num_output)
        response = self.tokenizer.decode(outputs[0])
        for token in response:
            yield CompletionResponse(text=token, delta=token)

offload_dir = r"F:\llm\offload"
if not os.path.exists(offload_dir):
    os.makedirs(offload_dir)

tokenizer = AutoTokenizer.from_pretrained("Qwen/Qwen2.5-1.5B-Instruct")
finetuned_model_path = r"F:\llm\Qwen2.5-1.5B-SFT+RAG\baseline\Qwen_QLoRA"

llm = CustomLLM(
    pretrained_model_name_or_path="Qwen/Qwen2.5-1.5B-Instruct",
    finetuned_model_path=finetuned_model_path,
    offload_dir=offload_dir
)

embedding_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")


Settings.llm = llm
Settings.embed_model = embedding_model

# 加载 .docx 文件
def load_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

# 将加载的文档内容转换为节点
file_path = "./data/1.docx"
doc_content = load_docx(file_path)

# 使用 SentenceSplitter 将文档分割成节点（RAG 需要节点形式的上下文）
node_parser = SentenceSplitter()

# 修改: 将字符串内容转换为 Document 对象
doc = LlamaDocument(text=doc_content)
nodes = node_parser.get_nodes_from_documents([doc])

# 构建向量索引（RAG 的基础）
vector_index = VectorStoreIndex(nodes)


# 构建向量索引和查询引擎
print(f'检索中...')
vector_index = VectorStoreIndex(nodes)
engine = vector_index.as_query_engine()

# 测试查询
question = "What can we gain from studying at SJTU?"  
response = engine.query(question)
answer = str(response)

# 打印结果
print(f"question={question}")
print(f'************')
print(f"Answer: {answer}")

# 评估答案相关性
evaluator = AnswerRelevancyEvaluator(llm)
result = evaluator.evaluate(query=question, response=answer)
print(f"score: {result.score}")
print(f"feedback: {result.feedback}")

# 评估上下文相关性
print(f'*****----------******')
contexts = [n.get_content() for n in response.source_nodes]
evaluator = ContextRelevancyEvaluator(llm)
result = evaluator.evaluate(query=question, contexts=contexts)
print(f"ContextRelevancy_score: {result.score}")
print(f"ContextRelevancy_feedback: {result.feedback}")
print(f'ContextRelevancy_query: {result.query}"')